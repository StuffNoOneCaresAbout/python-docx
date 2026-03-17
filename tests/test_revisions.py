from __future__ import annotations

from io import BytesIO

import pytest

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.revisions import TrackedDeletion, TrackedInsertion


def _revision_attrs(revision_id: int = 1, author: str = "TestAuthor") -> dict[str, str]:
    return {
        qn("w:id"): str(revision_id),
        qn("w:author"): author,
        qn("w:date"): "2026-03-17T00:00:00Z",
    }


def _insert_body_child(document: Document, element) -> None:
    body = document._element.body
    insert_idx = len(body)
    if len(body) and body[-1].tag == qn("w:sectPr"):
        insert_idx -= 1
    body.insert(insert_idx, element)


def _new_block_change(change_tag: str, text: str) -> object:
    change = OxmlElement(change_tag, attrs=_revision_attrs())
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_elm = OxmlElement("w:delText" if change_tag == "w:del" else "w:t")
    text_elm.text = text
    run.append(text_elm)
    paragraph.append(run)
    change.append(paragraph)
    return change


class DescribeRevisions:
    def it_excludes_inserted_text_from_paragraph_text_but_includes_it_in_accepted_text(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        insertion = paragraph.add_tracked_insertion("Beta", author="TestAuthor")

        assert isinstance(insertion, TrackedInsertion)
        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "AlphaBeta"
        assert paragraph.deleted_text == ""

    def it_includes_deleted_text_in_paragraph_text_but_not_in_accepted_text(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        deletion = paragraph.add_tracked_deletion(1, 4, author="TestAuthor")

        assert isinstance(deletion, TrackedDeletion)
        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "Aa"
        assert paragraph.deleted_text == "lph"

    def it_can_iterate_revisions_when_requested(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        paragraph.add_tracked_insertion("Beta", author="TestAuthor")
        paragraph.add_tracked_deletion(1, 3, author="TestAuthor")

        items = list(paragraph.iter_inner_content(include_revisions=True))

        assert any(isinstance(item, TrackedInsertion) for item in items)
        assert any(isinstance(item, TrackedDeletion) for item in items)

    def it_can_accept_and_reject_tracked_changes(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        insertion = paragraph.add_tracked_insertion("Beta", author="TestAuthor")
        deletion = paragraph.add_tracked_deletion(1, 4, author="TestAuthor")

        assert deletion is not None
        deletion.reject()
        insertion.reject()

        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "Alpha"
        assert paragraph.track_changes == []

    def it_can_find_and_replace_with_tracking(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        count = document.find_and_replace_tracked("ph", "XY", author="TestAuthor")

        assert count == 1
        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "AlXYa"
        assert len(document.track_changes) == 2

    def it_can_find_and_replace_with_tracking_in_nested_tables(self):
        document = Document()
        outer_table = document.add_table(rows=1, cols=1)
        inner_table = outer_table.cell(0, 0).add_table(rows=1, cols=1)
        paragraph = inner_table.cell(0, 0).paragraphs[0]
        paragraph.text = "Alpha"

        count = document.find_and_replace_tracked("ph", "XY", author="TestAuthor")

        assert count == 1
        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "AlXYa"

    def it_does_not_search_headers_or_footers_unless_requested(self):
        document = Document()
        paragraph = document.sections[0].header.paragraphs[0]
        paragraph.text = "Alpha"

        count = document.find_and_replace_tracked("ph", "XY", author="TestAuthor")

        assert count == 0
        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "Alpha"

    def it_can_optionally_find_and_replace_with_tracking_in_headers_and_footers(self):
        document = Document()
        header_paragraph = document.sections[0].header.paragraphs[0]
        footer_paragraph = document.sections[0].footer.paragraphs[0]
        header_paragraph.text = "Alpha"
        footer_paragraph.text = "Graph"

        count = document.find_and_replace_tracked(
            "ph",
            "XY",
            author="TestAuthor",
            include_headers_footers=True,
        )

        assert count == 2
        assert header_paragraph.text == "Alpha"
        assert header_paragraph.accepted_text == "AlXYa"
        assert footer_paragraph.text == "Graph"
        assert footer_paragraph.accepted_text == "GraXY"

    def it_does_not_create_header_or_footer_parts_when_optionally_searching_them(self):
        document = Document()
        section = document.sections[0]
        sectPr_xml_before = section._sectPr.xml  # pyright: ignore[reportPrivateUsage]

        count = document.find_and_replace_tracked(
            "ph",
            "XY",
            author="TestAuthor",
            include_headers_footers=True,
        )

        assert count == 0
        assert section._sectPr.xml == sectPr_xml_before  # pyright: ignore[reportPrivateUsage]

    def it_finds_matches_inside_existing_insertions(self):
        document = Document()
        paragraph = document.add_paragraph("Hello ")
        paragraph.add_tracked_insertion("World", author="TestAuthor")

        count = paragraph.replace_tracked("World", "Universe", author="TestAuthor")

        assert count == 1
        assert paragraph.accepted_text == "Hello Universe"
        assert [insertion.text for insertion in paragraph.insertions] == ["Universe"]
        assert len(paragraph.deletions) == 0

    def it_ignores_matches_that_exist_only_in_deleted_text(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")
        paragraph.add_tracked_deletion(0, 5, author="TestAuthor")

        count = paragraph.replace_tracked("Hello", "Hi", author="TestAuthor")

        assert count == 0
        assert paragraph.accepted_text == " World"
        assert paragraph.text == "Hello World"

    def it_replaces_text_spanning_multiple_runs(self):
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("ACME Corp")
        paragraph.add_run(" Ltd.")

        count = paragraph.replace_tracked("ACME Corp Ltd.", "ACME Inc Ltd.", author="TestAuthor")

        assert count == 1
        assert any(deletion.text == "ACME Corp Ltd." for deletion in paragraph.deletions)
        assert any(insertion.text == "ACME Inc Ltd." for insertion in paragraph.insertions)

    def it_preserves_surrounding_text_for_cross_run_replacement(self):
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("amount of 26.000")
        paragraph.add_run(" Euros (TWENTY")

        count = paragraph.replace_tracked("26.000 Euros", "30.000 Euros", author="TestAuthor")

        assert count == 1
        assert any(deletion.text == "26.000 Euros" for deletion in paragraph.deletions)
        assert any(insertion.text == "30.000 Euros" for insertion in paragraph.insertions)
        assert "amount of" in paragraph.accepted_text
        assert "(TWENTY" in paragraph.accepted_text

    def it_returns_track_changes_in_document_order(self):
        document = Document()
        paragraph = document.add_paragraph("colour")
        paragraph.add_tracked_deletion(0, 6, author="Editor")
        paragraph.add_tracked_insertion("color", author="Editor")

        changes = paragraph.track_changes

        assert len(changes) == 2
        assert isinstance(changes[0], TrackedDeletion)
        assert isinstance(changes[1], TrackedInsertion)

    def it_can_add_a_tracked_insertion_at_the_start_of_visible_text(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        insertion = paragraph.add_tracked_insertion_at(0, "X", author="TestAuthor")

        assert insertion.text == "X"
        assert paragraph.accepted_text == "XAlpha"
        assert paragraph.text == "Alpha"

    def it_can_add_a_tracked_insertion_in_the_middle_of_visible_text(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        insertion = paragraph.add_tracked_insertion_at(2, "X", author="TestAuthor")

        assert insertion.text == "X"
        assert paragraph.accepted_text == "AlXpha"
        assert paragraph.text == "Alpha"

    def it_can_add_a_tracked_insertion_at_the_end_of_visible_text(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        insertion = paragraph.add_tracked_insertion_at(
            len(paragraph.accepted_text),
            "X",
            author="TestAuthor",
        )

        assert insertion.text == "X"
        assert paragraph.accepted_text == "AlphaX"
        assert paragraph.text == "Alpha"

    def it_can_add_a_tracked_insertion_inside_existing_visible_content(self):
        document = Document()
        paragraph = document.add_paragraph("Hello ")
        paragraph.add_tracked_insertion("World", author="TestAuthor", revision_id=1)

        insertion = paragraph.add_tracked_insertion_at(8, "NEW", author="TestAuthor")

        assert insertion.text == "NEW"
        assert paragraph.accepted_text == "Hello WoNEWrld"
        assert any(existing.text == "NEW" for existing in paragraph.insertions)

    def it_uses_accepted_text_offsets_when_inserting_with_deletions_present(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")
        paragraph.add_tracked_deletion(0, 5, author="TestAuthor", revision_id=1)

        insertion = paragraph.add_tracked_insertion_at(1, "Brave", author="TestAuthor")

        assert insertion.text == "Brave"
        assert paragraph.accepted_text == " BraveWorld"
        assert paragraph.text == "Hello World"

    def it_can_add_a_tracked_insertion_before_a_unique_match(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")

        insertion = paragraph.add_tracked_insertion_before("World", "Brave ", author="TestAuthor")

        assert insertion.text == "Brave "
        assert paragraph.accepted_text == "Hello Brave World"

    def it_can_add_a_tracked_insertion_after_a_unique_match(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")

        insertion = paragraph.add_tracked_insertion_after("Hello", ", dear", author="TestAuthor")

        assert insertion.text == ", dear"
        assert paragraph.accepted_text == "Hello, dear World"

    def it_can_add_a_tracked_insertion_before_text_spanning_normal_and_inserted_content(self):
        document = Document()
        paragraph = document.add_paragraph("Hello ")
        paragraph.add_tracked_insertion("Brave ", author="TestAuthor", revision_id=1)
        paragraph.add_run("World")

        insertion = paragraph.add_tracked_insertion_before(
            "Brave World", "Dear ", author="TestAuthor"
        )

        assert insertion.text == "Dear "
        assert paragraph.accepted_text == "Hello Dear Brave World"

    def it_rejects_add_tracked_insertion_before_when_search_text_is_missing(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")

        with pytest.raises(ValueError, match="not found"):
            paragraph.add_tracked_insertion_before("Nope", "X", author="TestAuthor")

    def it_rejects_add_tracked_insertion_after_when_search_text_is_ambiguous(self):
        document = Document()
        paragraph = document.add_paragraph("World and World again")

        with pytest.raises(ValueError, match="multiple"):
            paragraph.add_tracked_insertion_after("World", "X", author="TestAuthor")

    def it_ignores_deleted_only_text_when_searching_for_insertion_position(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")
        paragraph.add_tracked_deletion(0, 5, author="TestAuthor", revision_id=1)

        with pytest.raises(ValueError, match="not found"):
            paragraph.add_tracked_insertion_before("Hello", "X", author="TestAuthor")

    def it_can_reject_all_tracked_changes_in_the_document(self):
        document = Document()
        first = document.add_paragraph("Alpha")
        second = document.add_paragraph("Gamma")
        first.add_tracked_insertion("Beta", author="TestAuthor")
        second.add_tracked_deletion(1, 3, author="TestAuthor")

        document.reject_all()

        assert document.track_changes == []
        assert first.text == "Alpha"
        assert first.accepted_text == "Alpha"
        assert second.text == "Gamma"
        assert second.accepted_text == "Gamma"

    def it_replaces_at_offsets_using_accepted_text_when_deletions_are_present(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")
        paragraph.add_tracked_deletion(0, 5, author="TestAuthor", revision_id=1)

        paragraph.replace_tracked_at(1, 6, "Universe", author="TestAuthor")

        assert paragraph.accepted_text == " Universe"
        assert [deletion.text for deletion in paragraph.deletions] == ["Hello", "World"]
        assert any(insertion.text == "Universe" for insertion in paragraph.insertions)

    def it_replaces_across_inserted_and_normal_visible_text(self):
        document = Document()
        paragraph = document.add_paragraph("Hello ")
        paragraph.add_tracked_insertion("Brave ", author="TestAuthor", revision_id=1)
        paragraph.add_run("World")

        paragraph.replace_tracked_at(
            6,
            len(paragraph.accepted_text),
            "Universe",
            author="TestAuthor",
        )

        assert paragraph.accepted_text == "Hello Universe"
        assert any(deletion.text == "World" for deletion in paragraph.deletions)
        assert any(insertion.text == "Universe" for insertion in paragraph.insertions)

    def it_can_replace_within_a_single_run_by_offset(self):
        document = Document()
        paragraph = document.add_paragraph("Hello World")

        paragraph.replace_tracked_at(6, 11, "Universe", author="TestAuthor")

        assert any(deletion.text == "World" for deletion in paragraph.deletions)
        assert any(insertion.text == "Universe" for insertion in paragraph.insertions)

    def it_can_replace_across_multiple_runs_by_offset(self):
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("Hello ")
        paragraph.add_run("World")

        paragraph.replace_tracked_at(4, 9, "X", author="TestAuthor")

        assert any(deletion.text == "o Wor" for deletion in paragraph.deletions)
        assert any(insertion.text == "X" for insertion in paragraph.insertions)
        assert "Hell" in paragraph.accepted_text
        assert "ld" in paragraph.accepted_text

    def it_can_replace_text_in_a_run_by_offset(self):
        document = Document()
        run = document.add_paragraph("Hello World").runs[0]

        run.replace_tracked_at(6, 11, "Universe", author="TestAuthor")

        paragraph = document.paragraphs[0]
        assert any(deletion.text == "World" for deletion in paragraph.deletions)
        assert any(insertion.text == "Universe" for insertion in paragraph.insertions)

    def it_removes_inserted_text_when_a_deletion_targets_only_that_visible_span(self):
        document = Document()
        paragraph = document.add_paragraph("Hello ")
        paragraph.add_tracked_insertion("World", author="TestAuthor", revision_id=1)

        tracked = paragraph.add_tracked_deletion(6, 11, author="TestAuthor")

        assert tracked is None
        assert paragraph.accepted_text == "Hello "
        assert paragraph.text == "Hello "
        assert len(paragraph.insertions) == 0

    def it_ignores_comment_markers_when_computing_paragraph_text_views(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        paragraph.add_comment("Comment", author="TestAuthor")
        paragraph.add_tracked_insertion("Beta", author="TestAuthor")
        paragraph.add_tracked_deletion(1, 4, author="TestAuthor")

        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "AaBeta"
        assert paragraph.deleted_text == "lph"

    def it_preserves_revision_text_semantics_on_save_and_reload(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        paragraph.add_tracked_insertion("Beta", author="TestAuthor")
        paragraph.add_tracked_deletion(1, 4, author="TestAuthor")

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        reloaded = Document(stream)
        paragraph = reloaded.paragraphs[0]

        assert paragraph.text == "Alpha"
        assert paragraph.accepted_text == "AaBeta"
        assert paragraph.deleted_text == "lph"

    def it_can_add_a_comment_to_a_run_level_tracked_insertion(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        insertion = paragraph.add_tracked_insertion("Beta", author="Editor")

        comment = insertion.add_comment("Review insertion", author="Reviewer")

        children = list(paragraph._p)
        assert comment.text == "Review insertion"
        assert [child.tag for child in children] == [
            qn("w:r"),
            qn("w:commentRangeStart"),
            qn("w:ins"),
            qn("w:commentRangeEnd"),
            qn("w:r"),
        ]
        assert children[1].get(qn("w:id")) == str(comment.comment_id)
        assert children[3].get(qn("w:id")) == str(comment.comment_id)
        assert children[4].xpath("./w:commentReference/@w:id") == [str(comment.comment_id)]

    def it_can_add_a_comment_to_a_run_level_tracked_deletion(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        deletion = paragraph.add_tracked_deletion(1, 4, author="Editor")

        comment = deletion.add_comment("Review deletion", author="Reviewer")

        children = list(paragraph._p)
        assert comment.text == "Review deletion"
        assert [child.tag for child in children] == [
            qn("w:r"),
            qn("w:commentRangeStart"),
            qn("w:del"),
            qn("w:commentRangeEnd"),
            qn("w:r"),
            qn("w:r"),
        ]
        assert children[1].get(qn("w:id")) == str(comment.comment_id)
        assert children[3].get(qn("w:id")) == str(comment.comment_id)
        assert children[4].xpath("./w:commentReference/@w:id") == [str(comment.comment_id)]

    def it_can_add_multiple_comments_to_the_same_tracked_change_span(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        insertion = paragraph.add_tracked_insertion("Beta", author="Editor")

        first = insertion.add_comment("First", author="Reviewer")
        second = insertion.add_comment("Second", author="Reviewer")

        children = list(paragraph._p)
        assert [child.tag for child in children] == [
            qn("w:r"),
            qn("w:commentRangeStart"),
            qn("w:commentRangeStart"),
            qn("w:ins"),
            qn("w:commentRangeEnd"),
            qn("w:r"),
            qn("w:commentRangeEnd"),
            qn("w:r"),
        ]
        assert children[1].get(qn("w:id")) == str(first.comment_id)
        assert children[2].get(qn("w:id")) == str(second.comment_id)
        assert children[4].get(qn("w:id")) == str(first.comment_id)
        assert children[6].get(qn("w:id")) == str(second.comment_id)
        assert children[5].xpath("./w:commentReference/@w:id") == [str(first.comment_id)]
        assert children[7].xpath("./w:commentReference/@w:id") == [str(second.comment_id)]

    def it_can_add_a_comment_to_a_block_level_tracked_paragraph(self):
        document = Document()
        insertion_elm = _new_block_change("w:ins", "Inserted paragraph")
        _insert_body_child(document, insertion_elm)
        insertion = TrackedInsertion(insertion_elm, document)

        comment = insertion.add_comment("Review block insertion", author="Reviewer")

        paragraph = insertion_elm.xpath("./w:p")[0]
        children = list(paragraph)
        assert comment.text == "Review block insertion"
        assert [child.tag for child in children] == [
            qn("w:commentRangeStart"),
            qn("w:r"),
            qn("w:commentRangeEnd"),
            qn("w:r"),
        ]
        assert insertion_elm.getparent()[0].tag == qn("w:ins")

    def it_can_add_a_comment_to_a_block_level_tracked_table(self):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell text"
        tbl = table._tbl
        parent = tbl.getparent()
        index = list(parent).index(tbl)
        parent.remove(tbl)
        deletion_elm = OxmlElement("w:del", attrs=_revision_attrs())
        deletion_elm.append(tbl)
        parent.insert(index, deletion_elm)
        deletion = TrackedDeletion(deletion_elm, document)

        comment = deletion.add_comment("Review table deletion", author="Reviewer")

        first_paragraph = deletion_elm.xpath(".//w:tc[1]/w:p[1]")[0]
        children = list(first_paragraph)
        assert comment.text == "Review table deletion"
        assert children[0].tag == qn("w:commentRangeStart")
        assert children[1].tag == qn("w:r")
        assert deletion_elm.xpath("./w:commentRangeStart") == []
        assert deletion_elm.xpath("./w:commentRangeEnd") == []

    def it_preserves_tracked_change_comments_on_save_and_reload(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        insertion = paragraph.add_tracked_insertion("Beta", author="Editor")
        comment = insertion.add_comment("Review insertion", author="Reviewer")

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        reloaded = Document(stream)
        paragraph = reloaded.paragraphs[0]
        reloaded_comment = next(iter(reloaded.comments))

        assert reloaded_comment.comment_id == comment.comment_id
        assert reloaded_comment.text == "Review insertion"
        assert paragraph._p.xpath("./w:commentRangeStart/@w:id") == [str(comment.comment_id)]
        assert paragraph._p.xpath("./w:commentRangeEnd/@w:id") == [str(comment.comment_id)]

    def it_rejects_tracked_change_comments_outside_the_main_document_story(self):
        document = Document()
        root_comment = document.add_paragraph("Alpha").add_comment("Root", author="Reviewer")
        insertion = root_comment.paragraphs[0].add_tracked_insertion("Beta", author="Editor")

        with pytest.raises(ValueError, match="main document story"):
            insertion.add_comment("Nested nope", author="Reviewer")

    def it_rejects_tracked_changes_with_no_anchorable_content(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        empty_insertion = OxmlElement("w:ins", attrs=_revision_attrs())
        paragraph._p.append(empty_insertion)
        insertion = TrackedInsertion(empty_insertion, paragraph)

        with pytest.raises(ValueError, match="no anchorable content"):
            insertion.add_comment("Nope", author="Reviewer")
