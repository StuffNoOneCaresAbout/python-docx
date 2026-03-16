# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.comments` module."""

from __future__ import annotations

import datetime as dt
from typing import cast

import pytest

from docx.comments import Comment, Comments
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.oxml.comments import (
    CT_Comment,
    CT_CommentEx,
    CT_Comments,
    CT_CommentsEx,
    CT_CommentsExtensible,
    CT_CommentsIds,
    CT_People,
)
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.package import Package
from docx.parts.comments import (
    CommentsExtendedPart,
    CommentsExtensiblePart,
    CommentsIdsPart,
    CommentsPart,
    PeoplePart,
)

from .unitutil.cxml import element, xml
from .unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeComments:
    """Unit-test suite for `docx.comments.Comments` objects."""

    @pytest.mark.parametrize(
        ("cxml", "count"),
        [
            ("w:comments", 0),
            ("w:comments/w:comment", 1),
            ("w:comments/(w:comment,w:comment,w:comment)", 3),
        ],
    )
    def it_knows_how_many_comments_it_contains(self, cxml: str, count: int, package_: Mock):
        comments_elm = cast(CT_Comments, element(cxml))
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        assert len(comments) == count

    def it_is_iterable_over_the_comments_it_contains(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments/(w:comment,w:comment)"))
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment_iter = iter(comments)

        comment1 = next(comment_iter)
        assert type(comment1) is Comment, "expected a `Comment` object"
        comment2 = next(comment_iter)
        assert type(comment2) is Comment, "expected a `Comment` object"
        with pytest.raises(StopIteration):
            next(comment_iter)

    def it_can_get_a_comment_by_id(self, package_: Mock):
        comments_elm = cast(
            CT_Comments,
            element("w:comments/(w:comment{w:id=1},w:comment{w:id=2},w:comment{w:id=3})"),
        )
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment = comments.get(2)

        assert type(comment) is Comment, "expected a `Comment` object"
        assert comment._comment_elm is comments_elm.comment_lst[1]

    def but_it_returns_None_when_no_comment_with_that_id_exists(self, package_: Mock):
        comments_elm = cast(
            CT_Comments,
            element("w:comments/(w:comment{w:id=1},w:comment{w:id=2},w:comment{w:id=3})"),
        )
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment = comments.get(4)

        assert comment is None, "expected None when no comment with that id exists"

    def it_can_add_a_new_comment(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        now_before = dt.datetime.now(dt.timezone.utc).replace(microsecond=0)
        comments = Comments(comments_elm, comments_part)

        comment = comments.add_comment()

        now_after = dt.datetime.now(dt.timezone.utc).replace(microsecond=0)
        # -- a comment is unconditionally added, and returned for any further adjustment --
        assert isinstance(comment, Comment)
        # -- it is "linked" to the comments part so it can add images and hyperlinks, etc. --
        assert comment.part is comments_part
        # -- comment numbering starts at 0, and is incremented for each new comment --
        assert comment.comment_id == 0
        # -- author is a required attribut, but is the empty string by default --
        assert comment.author == ""
        # -- initials is an optional attribute, but defaults to the empty string, same as Word --
        assert comment.initials == ""
        # -- timestamp is also optional, but defaults to now-UTC --
        assert comment.timestamp is not None
        assert now_before <= comment.timestamp <= now_after
        # -- by default, a new comment contains a single empty paragraph --
        assert [p.text for p in comment.paragraphs] == [""]
        # -- that paragraph has the "CommentText" style, same as Word applies --
        comment_elm = comment._comment_elm
        assert len(comment_elm.p_lst) == 1
        p = comment_elm.p_lst[0]
        assert p.style == "CommentText"
        assert p.paraId == comment.para_id == "00000001"
        assert p.textId == "77777777"
        # -- and that paragraph contains a single run with the necessary annotation reference --
        assert len(p.r_lst) == 1
        r = comment_elm.p_lst[0].r_lst[0]
        assert r.style == "CommentReference"
        assert r[-1].tag == qn("w:annotationRef")

    def it_does_not_create_threaded_reply_parts_for_a_top_level_comment(self, package_: Mock):
        comments, _ = _comments_with_extensions(package_)

        comments.add_comment("Parent")

        assert package_.main_document_part._comments_ids_part.element.commentId_lst == []
        assert (
            package_.main_document_part._comments_extensible_part.element.commentExtensible_lst
            == []
        )
        assert package_.main_document_part._people_part.element.person_lst == []

    def and_it_can_add_text_to_the_comment_when_adding_it(self, comments: Comments, package_: Mock):
        comment = comments.add_comment(text="para 1\n\npara 2")

        assert len(comment.paragraphs) == 3
        assert [p.text for p in comment.paragraphs] == ["para 1", "", "para 2"]
        assert all(p._p.style == "CommentText" for p in comment.paragraphs)

    def and_it_sets_the_author_and_their_initials_when_adding_a_comment_when_provided(
        self, comments: Comments, package_: Mock
    ):
        comment = comments.add_comment(author="Steve Canny", initials="SJC")

        assert comment.author == "Steve Canny"
        assert comment.initials == "SJC"

    def it_associates_extension_metadata_with_comments_when_iterating(self, package_: Mock):
        comments, _ = _comments_with_extensions(
            package_,
            "w:comments/(w:comment{w:id=1},w:comment{w:id=2})",
            (
                '<w15:commentEx w15:paraId="00000001"/>'
                '<w15:commentEx w15:paraId="00000002" w15:done="1"/>'
            ),
        )

        comments_ = list(comments)

        assert [comment.para_id for comment in comments_] == ["00000001", "00000002"]
        assert comments_[1].resolved is True

    def it_can_add_a_reply_comment(self, package_: Mock):
        comments, _ = _comments_with_extensions(package_)
        parent = comments.add_comment("Parent")

        reply = parent.add_reply("Reply", author="Steve", initials="SC")

        assert reply.text == "Reply"
        assert reply.author == "Steve"
        assert reply.initials == "SC"
        assert reply.parent_comment is not None
        assert reply.parent_comment.comment_id == parent.comment_id
        assert reply.parent_para_id == parent.para_id
        assert [comment.comment_id for comment in parent.replies] == [reply.comment_id]
        assert [
            comment_id.paraId
            for comment_id in package_.main_document_part._comments_ids_part.element.commentId_lst
        ] == [
            parent.para_id,
            reply.para_id,
        ]
        assert [
            comment_extensible.durableId
            for comment_extensible in package_.main_document_part._comments_extensible_part.element.commentExtensible_lst
        ] == [package_.main_document_part._comments_ids_part.element.commentId_lst[1].durableId]
        assert [
            person.author for person in package_.main_document_part._people_part.element.person_lst
        ] == ["Steve"]
        assert package_.main_document_part.element.xml == xml(
            "w:document/w:body/w:p/("
            "w:commentRangeStart{w:id=0},"
            "w:commentRangeStart{w:id=1},"
            'w:r/w:t"anchor",'
            "w:commentRangeEnd{w:id=0},"
            "w:r/(w:rPr/w:rStyle{w:val=CommentReference},w:commentReference{w:id=0}),"
            "w:commentRangeEnd{w:id=1},"
            "w:r/(w:rPr/w:rStyle{w:val=CommentReference},w:commentReference{w:id=1})"
            ")"
        )

    def it_can_mark_a_comment_resolved(self, package_: Mock):
        comments, _ = _comments_with_extensions(package_)
        comment = comments.add_comment("Needs review")

        comment.resolved = True

        assert comment.resolved is True

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def comments(self, package_: Mock) -> Comments:
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        return Comments(comments_elm, comments_part)

    @pytest.fixture
    def package_(self, request: FixtureRequest):
        package_ = instance_mock(request, Package)
        package_.main_document_part._comments_extended_part = CommentsExtendedPart(
            PackURI("/word/commentsExtended.xml"),
            CT.WML_COMMENTS_EXTENDED,
            _comments_ex_elm(),
            package_,
        )
        package_.main_document_part._comments_ids_part = CommentsIdsPart(
            PackURI("/word/commentsIds.xml"),
            CT.WML_COMMENTS_IDS,
            _comments_ids_elm(),
            package_,
        )
        package_.main_document_part._comments_extensible_part = CommentsExtensiblePart(
            PackURI("/word/commentsExtensible.xml"),
            CT.WML_COMMENTS_EXTENSIBLE,
            _comments_extensible_elm(),
            package_,
        )
        package_.main_document_part._people_part = PeoplePart(
            PackURI("/word/people.xml"),
            CT.WML_PEOPLE,
            _people_elm(),
            package_,
        )
        return package_


class DescribeComment:
    """Unit-test suite for `docx.comments.Comment`."""

    def it_knows_its_comment_id(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.comment_id == 42

    def it_knows_its_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:author=Steve Canny}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.author == "Steve Canny"

    def it_knows_the_initials_of_its_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:initials=SJC}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.initials == "SJC"

    def it_knows_the_date_and_time_it_was_authored(self, comments_part_: Mock):
        comment_elm = cast(
            CT_Comment,
            element("w:comment{w:id=42,w:date=2023-10-01T12:34:56Z}"),
        )
        comment = Comment(comment_elm, comments_part_)

        assert comment.timestamp == dt.datetime(2023, 10, 1, 12, 34, 56, tzinfo=dt.timezone.utc)

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:comment{w:id=42}", ""),
            ('w:comment{w:id=42}/w:p/w:r/w:t"Comment text."', "Comment text."),
            (
                'w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")',
                "First para\nSecond para",
            ),
            (
                'w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p,w:p/w:r/w:t"Second para")',
                "First para\n\nSecond para",
            ),
        ],
    )
    def it_can_summarize_its_content_as_text(
        self, cxml: str, expected_value: str, comments_part_: Mock
    ):
        assert Comment(cast(CT_Comment, element(cxml)), comments_part_).text == expected_value

    def it_provides_access_to_the_paragraphs_it_contains(self, comments_part_: Mock):
        comment_elm = cast(
            CT_Comment,
            element('w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")'),
        )
        comment = Comment(comment_elm, comments_part_)

        paragraphs = comment.paragraphs

        assert len(paragraphs) == 2
        assert [para.text for para in paragraphs] == ["First para", "Second para"]

    def it_can_update_the_comment_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:author=Old Author}"))
        comment = Comment(comment_elm, comments_part_)

        comment.author = "New Author"

        assert comment.author == "New Author"

    @pytest.mark.parametrize(
        "initials",
        [
            # -- valid initials --
            "XYZ",
            # -- empty string is valid
            "",
            # -- None is valid, removes existing initials
            None,
        ],
    )
    def it_can_update_the_comment_initials(self, initials: str | None, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:initials=ABC}"))
        comment = Comment(comment_elm, comments_part_)

        comment.initials = initials

        assert comment.initials == initials

    def it_knows_it_has_no_parent_when_not_a_reply(self, comments_part_: Mock):
        comment = Comment(cast(CT_Comment, element("w:comment{w:id=42}")), comments_part_)

        assert comment.parent_comment is None
        assert comment.parent_para_id is None
        assert comment.para_id is None
        assert comment.replies == []
        assert comment.resolved is False

    def it_can_create_extension_metadata_when_marked_resolved(
        self, comments_part_: Mock, comment_ex_: Mock
    ):
        comment = Comment(cast(CT_Comment, element("w:comment{w:id=42}")), comments_part_)
        comments_part_.ensure_comment_ex.return_value = comment_ex_

        comment.resolved = True

        comments_part_.ensure_comment_ex.assert_called_once_with(comment._comment_elm)
        assert comment_ex_.done is True
        assert comment.resolved is True

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def comments_part_(self, request: FixtureRequest):
        return instance_mock(request, CommentsPart)

    @pytest.fixture
    def comment_ex_(self, request: FixtureRequest):
        return instance_mock(request, CT_CommentEx)


def _comments_with_extensions(
    package_: Mock,
    comments_cxml: str = "w:comments",
    comments_ex_xml: str = "",
) -> tuple[Comments, CommentsPart]:
    comments_elm = cast(CT_Comments, element(comments_cxml))
    comments_part = CommentsPart(
        PackURI("/word/comments.xml"), CT.WML_COMMENTS, comments_elm, package_
    )
    comments_extended_part = CommentsExtendedPart(
        PackURI("/word/commentsExtended.xml"),
        CT.WML_COMMENTS_EXTENDED,
        _comments_ex_elm(comments_ex_xml),
        package_,
    )
    package_.main_document_part.element = cast(
        object,
        element(
            "w:document/w:body/w:p/("
            "w:commentRangeStart{w:id=0},"
            'w:r/w:t"anchor",'
            "w:commentRangeEnd{w:id=0},"
            "w:r/(w:rPr/w:rStyle{w:val=CommentReference},w:commentReference{w:id=0})"
            ")"
        ),
    )
    package_.main_document_part._comments_extended_part = comments_extended_part
    return Comments(comments_elm, comments_part), comments_part


def _comments_ex_elm(inner_xml: str = "") -> CT_CommentsEx:
    return cast(
        CT_CommentsEx,
        parse_xml(
            '<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
            f"{inner_xml}"
            "</w15:commentsEx>"
        ),
    )


def _comments_ids_elm(inner_xml: str = "") -> CT_CommentsIds:
    return cast(
        CT_CommentsIds,
        parse_xml(
            '<w16cid:commentsIds xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid">'
            f"{inner_xml}"
            "</w16cid:commentsIds>"
        ),
    )


def _comments_extensible_elm(inner_xml: str = "") -> CT_CommentsExtensible:
    return cast(
        CT_CommentsExtensible,
        parse_xml(
            '<w16cex:commentsExtensible xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex">'
            f"{inner_xml}"
            "</w16cex:commentsExtensible>"
        ),
    )


def _people_elm(inner_xml: str = "") -> CT_People:
    return cast(
        CT_People,
        parse_xml(
            '<w15:people xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
            f"{inner_xml}"
            "</w15:people>"
        ),
    )
