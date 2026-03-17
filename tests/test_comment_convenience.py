from __future__ import annotations

import datetime as dt
from io import BytesIO
from zipfile import ZipFile

import pytest

from docx import Document


class DescribeCommentConvenience:
    def it_can_add_a_comment_from_a_paragraph(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        comment = paragraph.add_comment("Note", author="TestAuthor", initials="CP")

        assert len(document.comments) == 1
        assert comment.text == "Note"
        assert comment.author == "TestAuthor"
        assert comment.initials == "CP"

    def it_can_add_a_comment_from_a_run(self):
        document = Document()
        run = document.add_paragraph("Alpha").runs[0]

        comment = run.add_comment("Run note", author="TestAuthor")

        assert len(document.comments) == 1
        assert comment.text == "Run note"

    def it_can_add_a_comment_with_an_explicit_timestamp_from_a_paragraph(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        timestamp = dt.datetime(2025, 6, 11, 20, 42, 30, tzinfo=dt.timezone(dt.timedelta(hours=8)))

        comment = paragraph.add_comment("Note", author="TestAuthor", timestamp=timestamp)

        assert comment.timestamp == dt.datetime(2025, 6, 11, 12, 42, 30, tzinfo=dt.timezone.utc)

    def it_writes_word_style_timestamp_parts_for_timezone_aware_comments(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        timestamp = dt.datetime(2025, 6, 11, 20, 42, 30, tzinfo=dt.timezone(dt.timedelta(hours=8)))

        paragraph.add_comment("Note", author="TestAuthor", timestamp=timestamp)

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        with ZipFile(stream) as pkg:
            comments_xml = pkg.read("word/comments.xml").decode("utf-8")
            comments_ids_xml = pkg.read("word/commentsIds.xml").decode("utf-8")
            comments_extensible_xml = pkg.read("word/commentsExtensible.xml").decode("utf-8")

        assert 'w:date="2025-06-11T20:42:30"' in comments_xml
        assert "w16cid:commentId" in comments_ids_xml
        assert 'w16cex:dateUtc="2025-06-11T12:42:30' in comments_extensible_xml

    def it_stamps_resolution_timestamp_when_comment_is_resolved(self):
        document = Document()
        run = document.add_paragraph("Alpha").runs[0]
        comment = run.add_comment("Resolve me", author="TestAuthor")

        assert comment.resolved is False
        assert comment.resolved_at is None

        comment.resolve()

        assert comment.resolved is True
        assert comment.resolved_at is not None

    def it_can_stamp_an_explicit_resolution_timestamp(self):
        document = Document()
        run = document.add_paragraph("Alpha").runs[0]
        comment = run.add_comment("Resolve me", author="TestAuthor")
        timestamp = dt.datetime(2025, 6, 11, 20, 42, 30, tzinfo=dt.timezone(dt.timedelta(hours=8)))

        comment.resolve(timestamp=timestamp)

        assert comment.resolved_at == dt.datetime(2025, 6, 11, 12, 42, 30, tzinfo=dt.timezone.utc)

    def it_preserves_resolution_state_and_timestamp_on_save_and_reload(self):
        document = Document()
        run = document.add_paragraph("Alpha").runs[0]
        comment = run.add_comment("Resolve me", author="TestAuthor")
        comment.resolve()

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        reloaded = Document(stream)
        comment = next(iter(reloaded.comments))

        assert comment.resolved is True
        assert comment.resolved_at is not None

    def it_disallows_resolving_a_reply(self):
        document = Document()
        run = document.add_paragraph("Alpha").runs[0]
        parent = run.add_comment("Parent", author="TestAuthor")
        reply = parent.add_reply("Reply", author="Reviewer")

        with pytest.raises(ValueError, match="reply comments do not support resolved state"):
            reply.resolve()

        with pytest.raises(ValueError, match="reply comments do not support resolved state"):
            reply.reopen()

        with pytest.raises(ValueError, match="reply comments do not support resolved state"):
            reply.resolved = True

        assert reply.resolved is False
        assert reply.resolved_at is None

    def it_can_add_a_comment_to_a_substring_within_a_single_run(self):
        document = Document()
        paragraph = document.add_paragraph("South")

        comment = paragraph.add_comment_range(1, 3, "Substring note", author="TestAuthor")

        assert comment.text == "Substring note"
        assert paragraph.text == "South"
        assert [run.text for run in paragraph.runs] == ["S", "ou", "", "th"]

    def it_can_add_a_comment_to_a_substring_spanning_multiple_runs(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("So")
        paragraph.add_run("ut")
        paragraph.add_run("h")

        comment = paragraph.add_comment_range(1, 4, "Substring note", author="TestAuthor")

        assert comment.text == "Substring note"
        assert paragraph.text == "South"
        assert [run.text for run in paragraph.runs] == ["S", "o", "ut", "", "h"]

    def it_rejects_invalid_comment_range_offsets(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")

        with pytest.raises(ValueError, match="Invalid offsets"):
            paragraph.add_comment_range(2, 2, "Nope")

        with pytest.raises(ValueError, match="Invalid offsets"):
            paragraph.add_comment_range(-1, 2, "Nope")

        with pytest.raises(ValueError, match="Invalid offsets"):
            paragraph.add_comment_range(0, 6, "Nope")

    def it_rejects_comment_ranges_that_include_non_run_content(self):
        document = Document()
        paragraph = document.add_paragraph("Alpha")
        paragraph.add_tracked_insertion("Beta", author="Editor")

        with pytest.raises(
            ValueError,
            match="comment ranges currently support only plain paragraph runs",
        ):
            paragraph.add_comment_range(3, 7, "Nope")

    def it_preserves_substring_comments_on_save_and_reload(self):
        document = Document()
        paragraph = document.add_paragraph("South")
        paragraph.add_comment_range(1, 3, "Substring note", author="TestAuthor")

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        reloaded = Document(stream)
        paragraph = reloaded.paragraphs[0]
        comment = next(iter(reloaded.comments))

        assert comment.text == "Substring note"
        assert paragraph.text == "South"
